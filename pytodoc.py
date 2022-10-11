from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH ,WD_LINE_SPACING
import json
from docx.shared import RGBColor
from docx2pdf import convert


##https://www.figma.com/file/hJosoQ91EiZwQSCC0orUuJ/My-TEnS-FAB-Digital?node-id=87%3A1991 link figma baru
def _set_cell_background(cell, fill, color='black', val=None): #color table
    from docx.oxml.shared import qn  # feel free to move these out
    from docx.oxml.xmlchemy import OxmlElement

    cell_properties = cell._element.tcPr
    try:
        cell_shading = cell_properties.xpath('w:shd')[0]  # in case there's already shading
    except IndexError:
        cell_shading = OxmlElement('w:shd') # add new w:shd element to it
    if fill:
        cell_shading.set(qn('w:fill'), fill)  # set fill property, respecting namespace
    cell_properties.append(cell_shading) 

def get_val(x,data):  #function get val dict
    for key, value in data.items():
      if x == key:
        return value

def check_type(data):
  if isinstance(data,str):
    return 'string'
  elif isinstance(data,int):
    return 'number'
  else:
    return 'array/object'

document = Document()
section = document.sections
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

style = document.styles['Normal']
font = style.font
font.name ='Arial'
font.size = Pt(11)


true = 'true'
false = 'false'
null =''
with open('input.json','r', encoding="utf8") as json_file:
    jsonData = json.load(json_file)
path = jsonData['paths']
print(len(path)) 
components = jsonData['components']
tot = 0
for key in path:
  for method in path[key]:
    DescAPI = path[key][method]['summary']
    try:
      security = path[key][method]['security']
      sec = list(security[0].keys())[0]
    except:
      sec = ''
    table = document.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    #header
    table.rows[0].cells[0].merge(table.rows[0].cells[1])
    table.rows[0].cells[0].text = 'API URL - '+ DescAPI
    table.rows[1].cells[0].merge(table.rows[1].cells[1])
    table.rows[1].cells[0].text = '{{baseURL}}'+key
    table.rows[1].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    _set_cell_background(table.rows[1].cells[0], '666666')
    #top
    paragraph = document.add_paragraph()
    paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    table1 = document.add_table(rows=4, cols=2)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    _set_cell_background(table1.rows[0].cells[0], 'D9D9D9')
    _set_cell_background(table1.rows[1].cells[0], 'D9D9D9')
    _set_cell_background(table1.rows[2].cells[0], 'D9D9D9')
    _set_cell_background(table1.rows[3].cells[0], 'D9D9D9')
    hdr_cells[0].text = 'Description'
    hdr_cells[1].text = DescAPI
    hdr_cells = table1.rows[1].cells
    hdr_cells[0].text = 'Method'
    hdr_cells[1].text = method
    hdr_cells = table1.rows[2].cells
    hdr_cells[0].text = 'Content-Type'
    hdr_cells[1].text = 'application/json'
    hdr_cells = table1.rows[3].cells
    hdr_cells[0].text = 'Header'
    hdr_cells[1].text = sec
    
    try:
      i = len(path[key][method]['parameters'])
    except:
      i = 0
    if i > 0 : #paramaters doc
      paragraph = document.add_paragraph()
      paragraph = document.add_paragraph('Parameters :')
      tablea = document.add_table(rows=1, cols=4)
      bod = tablea.rows[0].cells
      bod[0].text = 'Attribute Name'
      bod[1].text = 'Data Type'
      bod[2].text = 'Description'
      bod[3].text = 'Example'
      _set_cell_background(tablea.rows[0].cells[0], 'c9211e')
      _set_cell_background(tablea.rows[0].cells[1], 'c9211e')
      _set_cell_background(tablea.rows[0].cells[2], 'c9211e')
      _set_cell_background(tablea.rows[0].cells[3], 'c9211e')
      bod[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
      bod[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
      bod[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
      bod[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
      u=0
      tablea2 = document.add_table(rows=i, cols=4)
      for param in path[key][method]['parameters']:
        if '$ref' in param:
          px = str(param['$ref'])
          pxc = px.replace("#/components/parameters/","")
          schemas = components['parameters'][pxc]
          Pdesc = ''
          Ptype = ''
          Pexamp = ''
          if "description" in schemas:
            Pdesc = schemas['description']
          if "schema" in schemas:
            if 'type' in schemas['schema']:
              Ptype = schemas['schema']['type']
            if 'example' in schemas['schema']:
              Pexamp = schemas['schema']['example']
          if 'example' in schemas:
            Pexamp = schemas['example']
          body = tablea2.rows[u].cells
          body[0].text = schemas['name']
          body[1].text = Ptype
          body[2].text = Pdesc
          body[3].text = Pexamp
          u+=1
        else:
          Pdesc = ''
          Ptype = ''
          Pexamp = ''
          if "description" in param:
            Pdesc = param['description']
          if "schema" in param:
            if 'type' in param['schema']:
              Ptype = param['schema']['type']
            if 'example' in param['schema']:
              Pexamp = param['schema']['example']
          if 'example' in param:
            Pexamp = param['example']
          else:
            Pexamp=param
          body = tablea2.rows[u].cells
          body[0].text = param['name']
          body[1].text = Ptype
          body[2].text = Pdesc
          body[3].text = Pexamp
          u+=1
    try: #request body doc
      for typeBody,valueRequest in path[key][method]['requestBody']['content'].items():
        paragraph = document.add_paragraph()
        paragraph = document.add_paragraph('Body :')
        table2a = document.add_table(rows=1, cols=4)
        bod = table2a.rows[0].cells
        bod[0].text = 'Attribute Name'
        bod[1].text = 'Data Type'
        bod[2].text = 'Description'
        bod[3].text = 'Example'
        _set_cell_background(table2a.rows[0].cells[0], 'c9211e')
        _set_cell_background(table2a.rows[0].cells[1], 'c9211e')
        _set_cell_background(table2a.rows[0].cells[2], 'c9211e')
        _set_cell_background(table2a.rows[0].cells[3], 'c9211e')
        bod[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        bod[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        bod[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        bod[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        # if dic check dict schema dan example 
        if 'example' in path[key][method]['requestBody']['content'][typeBody].keys(): #bodyreq with example
          x = path[key][method]['requestBody']['content'][typeBody]['example']
          
          if (isinstance(x, list)) : 
            i = len(path[key][method]['requestBody']['content'][typeBody]['example'])
            u=0
            table2 = document.add_table(rows=i, cols=4)
            for valueBody in path[key][method]['requestBody']['content'][typeBody]['example']:
              body = table2.rows[u].cells
              body[0].text = valueBody['key']
              body[1].text = valueBody['type']
              body[2].text = ''
              body[3].text = json.dumps(valueBody['value'], indent=4, sort_keys=True)
              u+=1  
            #exampleBody
            paragraph = document.add_paragraph()
            paragraph = document.add_paragraph('Example Body :')
            table13 = document.add_table(rows=1, cols=1)
            table13.rows[0].cells[0].text = json.dumps(path[key][method]['requestBody']['content'][typeBody]['example'], indent=4, sort_keys=True)
          else:
            i = len(path[key][method]['requestBody']['content'][typeBody]['example'].items())
            u=0
            table2 = document.add_table(rows=i, cols=4)
            for keyBody,valueBody in path[key][method]['requestBody']['content'][typeBody]['example'].items():
              body = table2.rows[u].cells
              body[0].text = keyBody
              body[1].text = ''
              body[2].text = ''
              body[3].text = json.dumps(valueBody, indent=4, sort_keys=True)
              u+=1  
            #exampleBody
            paragraph = document.add_paragraph()
            paragraph = document.add_paragraph('Example Body :')
            table13 = document.add_table(rows=1, cols=1)
            table13.rows[0].cells[0].text = json.dumps(path[key][method]['requestBody']['content'][typeBody]['example'], indent=4, sort_keys=True)
        elif 'examples' in path[key][method]['requestBody']['content'][typeBody].keys(): #bodyreq with example
          i = len(path[key][method]['requestBody']['content'][typeBody]['examples'].items())
          u=0
          table2 = document.add_table(rows=i, cols=4)
          for keyBody,valueBody in path[key][method]['requestBody']['content'][typeBody]['examples'].items():
            body = table2.rows[u].cells
            body[0].text = keyBody
            body[1].text = str(check_type(valueBody['value']))
            body[2].text = 'DATA PERLU DI CEK'
            body[3].text = json.dumps(valueBody['value'], indent=4, sort_keys=True)
            u+=1  
          #exampleBody
          paragraph = document.add_paragraph()
          paragraph = document.add_paragraph('Example Body :')
          table13 = document.add_table(rows=1, cols=1)
          table13.rows[0].cells[0].text = json.dumps(path[key][method]['requestBody']['content'][typeBody]['examples'], indent=4, sort_keys=True)
        elif 'schema' in path[key][method]['requestBody']['content'][typeBody].keys(): #bodyreq with schemas 
          bx = str(path[key][method]['requestBody']['content'][typeBody]['schema'])
          try:
            bxc = bx.replace("{'$ref': '#/components/schemas/","")
            bcc = bxc.replace("'}","")
            schemasb = components['schemas'][bcc]['properties']
            datb = {}
            valPropertiesb = {}
            dvalPropertiesb ={}
            dvalPropertiesbdict ={}
            for keySchemasb,valScehmasb in schemasb.items():
              try:
                datb.update({keySchemasb: valScehmasb['example']})
              except: #schemas dict properties
                try:
                  for keyPropb,valPropb in schemasb[keySchemasb]['properties'].items():
                    try:
                      valPropertiesb.update({keyPropb: valPropb['example']})
                    except:
                      for dkPropb, dvPropb in valPropb['properties'].items():
                        try:
                          dvalPropertiesb.update({dkPropb: dvPropb['example']})
                        except:
                          for keydvPropb, valdvPropb in dvPropb['properties'].items():
                            dvalPropertiesbdict.update({keydvPropb: valdvPropb['example']})
                          datb.update({dkPropb: dvalPropertiesbdict})
                      datb.update({keyPropb: dvalPropertiesb})
                except:
                  datb.update({keySchemasb:'example belum diisi'})
            i = len(datb.items())
            u=0
            table2 = document.add_table(rows=i, cols=4)
            for keyBody,valueBody in datb.items():
              body = table2.rows[u].cells
              body[0].text = keyBody
              body[1].text = str(check_type(valueBody))
              body[2].text = ''
              body[3].text = json.dumps(valueBody, indent=4, sort_keys=True)
              u+=1
            #exampleBody
            paragraph = document.add_paragraph()
            paragraph = document.add_paragraph('Example Body :')
            table13 = document.add_table(rows=1, cols=1)
            table13.rows[0].cells[0].text = json.dumps(datb, indent=4, sort_keys=True)
          except:
            i = len(path[key][method]['requestBody']['content'][typeBody]['schema']['example'].items())
            u=0
            table2 = document.add_table(rows=i, cols=4)
            for keyBody,valueBody in path[key][method]['requestBody']['content'][typeBody]['schema']['example'].items():
              body = table2.rows[u].cells
              body[0].text = keyBody
              body[1].text = str(check_type(valueBody))
              body[2].text = ''
              body[3].text = json.dumps(valueBody, indent=4, sort_keys=True)
              u+=1  
            #exampleBody
            paragraph = document.add_paragraph()
            paragraph = document.add_paragraph('Example Body :')
            table13 = document.add_table(rows=1, cols=1)
            table13.rows[0].cells[0].text = json.dumps(path[key][method]['requestBody']['content'][typeBody]['schema']['example'], indent=4, sort_keys=True)
    except:
      next
    try: #response doc
      if 'responses' in path[key][method]:
        paragraph = document.add_paragraph()
        paragraph = document.add_paragraph('Response :')
        r=0
      for codeResp,valueResponse in path[key][method]['responses'].items():
        try:
          if r>0:
            paragraph = document.add_paragraph()
          paragraph = document.add_paragraph(codeResp)
          table4a = document.add_table(rows=1, cols=4)
          bod = table4a.rows[0].cells
          bod[0].text = 'Attribute Name'
          bod[1].text = 'Data Type'
          bod[2].text = 'Description'
          bod[3].text = 'Example'
          _set_cell_background(table4a.rows[0].cells[0], 'c9211e')
          _set_cell_background(table4a.rows[0].cells[1], 'c9211e')
          _set_cell_background(table4a.rows[0].cells[2], 'c9211e')
          _set_cell_background(table4a.rows[0].cells[3], 'c9211e')
          bod[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
          bod[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
          bod[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
          bod[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
          if 'example' in valueResponse['content']['application/json'].keys(): #response check example/schema
            if 'data' in valueResponse['content']['application/json']['example'].keys():
              try: #data response object
                u = 0
                try:
                  i = len(list(valueResponse['content']['application/json']['example']['data'].items()))
                  table4 = document.add_table(rows=i, cols=4) 
                  for keyResp,valueResp in valueResponse['content']['application/json']['example']['data'].items():
                    body = table4.rows[u].cells
                    body[0].text = keyResp
                    body[1].text = str(check_type(valueResp))
                    body[2].text = ''
                    body[3].text = json.dumps(valueResp, indent=4, sort_keys=True)
                    u+=1
                except:
                  i = len(list(valueResponse['content']['application/json']['example']['data'][0].items()))
                  table4 = document.add_table(rows=i, cols=4) 
                  for keyResp,valueResp in valueResponse['content']['application/json']['example']['data'][0].items():
                    body = table4.rows[u].cells
                    body[0].text = keyResp
                    body[1].text = str(check_type(valueResp))
                    body[2].text = ''
                    body[3].text = json.dumps(valueResp, indent=4, sort_keys=True)
                    u+=1
                #exampleResp
                paragraph = document.add_paragraph()
                paragraph = document.add_paragraph('Example Response :')
                table5 = document.add_table(rows=1, cols=1)
                table5.rows[0].cells[0].text = json.dumps(valueResponse['content']['application/json']['example'], indent=4, sort_keys=True)
              except: #data response string
                table4 = document.add_table(rows=1, cols=4)
                resp = get_val('data',valueResponse['content']['application/json']['example'])
                body = table4.rows[0].cells
                body[0].text = 'data'
                body[1].text = str(check_type(resp))
                body[2].text = ''
                body[3].text = resp 
                paragraph = document.add_paragraph()
                paragraph = document.add_paragraph('Example Response :')
                table5 = document.add_table(rows=1, cols=1)
                table5.rows[0].cells[0].text = json.dumps(valueResponse['content']['application/json']['example'], indent=4, sort_keys=True)
            else:
              u = 0
              i = len(list(valueResponse['content']['application/json']['example'].items()))
              table4 = document.add_table(rows=i, cols=4)
              for keyA,valueA in valueResponse['content']['application/json']['example'].items():
                body = table4.rows[u].cells
                body[0].text = keyA
                body[1].text = str(check_type(valueA))
                body[2].text = ''
                body[3].text = json.dumps(valueA, indent=4, sort_keys=True)
                u+=1
              paragraph = document.add_paragraph()
              paragraph = document.add_paragraph('Example Response :')
              table5 = document.add_table(rows=1, cols=1)
              table5.rows[0].cells[0].text = json.dumps(valueResponse['content']['application/json']['example'], indent=4, sort_keys=True)
          elif 'examples' in valueResponse['content']['application/json'].keys(): #response check example/schema
            try: #data response object
              i = len(list(valueResponse['content']['application/json']['examples']['data'].items()))
              table4 = document.add_table(rows=i, cols=4)
              u = 0
              for keyResp,valueResp in valueResponse['content']['application/json']['examples']['data'].items():
                body = table4.rows[u].cells
                body[0].text = keyResp
                body[1].text = str(check_type(valueResp))
                body[2].text = ''
                body[3].text = json.dumps(valueResp, indent=4, sort_keys=True)
                u+=1
              #exampleResp
              paragraph = document.add_paragraph()
              paragraph = document.add_paragraph('Example Response :')
              table5 = document.add_table(rows=1, cols=1)
              table5.rows[0].cells[0].text = json.dumps(valueResponse['content']['application/json']['examples'], indent=4, sort_keys=True)
            except: #data response string
              table4 = document.add_table(rows=1, cols=4)
              resp = get_val('data',valueResponse['content']['application/json']['examples'])
              body = table4.rows[0].cells
              body[0].text = 'data'
              body[1].text = str(check_type(resp))
              body[2].text = ''
              body[3].text = resp 
              paragraph = document.add_paragraph()
              paragraph = document.add_paragraph('Example Response :')
              table5 = document.add_table(rows=1, cols=1)
              table5.rows[0].cells[0].text = json.dumps(valueResponse['content']['application/json']['examples'], indent=4, sort_keys=True)
          elif 'schema' in valueResponse['content']['application/json'].keys(): #response use schema
            ax = str(valueResponse['content']['application/json']['schema'])
            axc = ax.replace("{'$ref': '#/components/schemas/","")
            acc = axc.replace("'}","")
            schemas = components['schemas'][acc]['properties']
            dat = {}
            valProperties = {}
            for keySchemas,valScehmas in schemas['data'].items():
              if keySchemas == 'example':
                valExam = valScehmas
                dat.update({'data': valExam})
              if keySchemas == 'properties':
                for keySchemasp,valScehmasp in valScehmas.items():
                  try:
                    dat.update({keySchemasp: valScehmasp['example']})
                  except:
                    try:
                      for keySchemaspd,valScehmaspd in valScehmasp['properties'].items():
                        try:
                          dat.update({keySchemaspd: valScehmaspd['example']})
                        except:
                          dat.update({keySchemaspd: 'apa isinya?'})
                    except:
                      dat.update({keySchemasp: 'apa isinya?'})
            try:
              i = len(list(dat.items()))
              table4 = document.add_table(rows=i, cols=4)
              u = 0
              for keyResp,valueResp in dat.items():
                body = table4.rows[u].cells
                body[0].text = keyResp
                body[1].text = str(check_type(valueResp))
                body[2].text = ''
                body[3].text = json.dumps(valueResp, indent=4, sort_keys=True)
                u+=1
              #exampleResp
              paragraph = document.add_paragraph()
              paragraph = document.add_paragraph('Example Response :')
              table5 = document.add_table(rows=1, cols=1)
              table5.rows[0].cells[0].text = json.dumps(schemas, indent=4, sort_keys=True)
            except:
              table4 = document.add_table(rows=1, cols=4)
              resp = get_val('data',dat)
              body = table4.rows[0].cells
              body[0].text = 'data'
              body[1].text = str(check_type(resp))
              body[2].text = ''
              body[3].text = resp 
              paragraph = document.add_paragraph()
              paragraph = document.add_paragraph('Example Response :')
              table5 = document.add_table(rows=1, cols=1)
              table5.rows[0].cells[0].text = json.dumps(dat, indent=4, sort_keys=True)
        except:
          next
        r+1
    except:
      next
    document.add_page_break()
    tot+=1
print(tot)
document.save("data.docx")